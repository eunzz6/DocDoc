from django.shortcuts import render, redirect
from .forms import PersonForm
from .models import Person, Submit
import speech_recognition as sr
from pydub import AudioSegment
from django.contrib.auth.decorators import login_required
from pathlib import Path
import json,requests
import joblib
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
from sklearn.preprocessing import LabelBinarizer
from docx import Document
from docx2pdf import convert
import os
import matplotlib.pyplot as plt
from django.views.decorators.cache import never_cache
import pandas as pd
import numpy as np
from konlpy.tag import Kkma




model = joblib.load(r'C:\Users\User\WEB\34\web\model\doc_model.pkl')
max_len=14
tokenizer=joblib.load(r"C:\Users\User\WEB\34\web\model\tokenizer.joblib")
label_binarizer = joblib.load(r"C:\Users\User\WEB\34\web\model\labelbinarizer.joblib")



@login_required(login_url='/')
@never_cache
@login_required(login_url='/')
def file_upload(request):
    form = PersonForm(request.POST, request.FILES)
    return render(request,'main.html',{'form':form})


@login_required(login_url='/')
@never_cache
def submit(request):
    if request.method == 'POST':
        wearable_data = request.FILES['wearableData']
        form = PersonForm(request.POST, request.FILES)
        if form.is_valid():
            name = form.cleaned_data['name']
            age = form.cleaned_data['age']
            # file = form.cleaned_data['file']
            file2 = form.cleaned_data['file2']
        file = wearable_data
        submit = Submit(file=file)
        submit.save()
        #--------------------------------주석 풀면 돈이야 돈-------------------------------------------
        print(submit.file.path)
        res = ClovaSpeechClient().req_upload(file=submit.file.path, completion='sync')
        result = json.loads(res.text)

        segments = result['segments']
        speakers = result['speakers']

        speaker_results = {}
        new_texts = []
        for segment in segments:
            speaker_label = segment['speaker']['label']
            speaker_text = segment['text']
            if speaker_label in speaker_results:
                speaker_results[speaker_label].append(speaker_text)
            else:
                speaker_results[speaker_label] = [speaker_text]

            
            if speaker_label == '2':  # Check if the speaker is labeled as '2'
                new_texts.append(speaker_text)

        for speaker_label, texts in speaker_results.items():
            print(f"발화자 {speaker_label}의 결과:")
            print("\n".join(texts))
            print() 
       #-------------------------------------------------------------------------------------------------
      

        def tokenize_sentences(text):
            kkma = Kkma()
            sentences = kkma.sentences(text)
            return sentences
        
        # new_texts=['갑자기 심장이 빨리 뛰어요.',
        # '공복 혈압도 많이 낮은 것 같아요.',
        # '빈혈 증상도 같이 오는 기분이에요. 그리고 배가 너무 아파요',
        # '소화가 너무 안 되는 느낌이 들어요.',
        # '입맛이 없어서 아무것도 안 먹고 있어요.',
        # '혹시 머리도 너무 아픈데 두통약도 처방해 주실 수 있나요.',
        # '속이 울렁거리고 식욕이 없어요. 그래서 살이 3kg가 빠졌어요.',
        # '가끔 어지러운데 심하진 않아요.',
        # '아니요. 똑같아요.',
        # '네 감사합니다.']

        tokenized_sentences = []
        
        for text in new_texts:
            sentences = tokenize_sentences(text)
            tokenized_sentences.extend(sentences)
        print(tokenized_sentences)
        
        new_sequences = tokenizer.texts_to_sequences(tokenized_sentences)
        new_padded_sequences = pad_sequences(new_sequences, maxlen=14)
        predictions = model.predict(new_padded_sequences)

        decoded_labels = label_binarizer.inverse_transform(predictions)
        
        
        label_names = label_binarizer.classes_
        top_k = 10 

        top_k_indices = np.argsort(-predictions, axis=1)[:, :top_k]
        top_k_probabilities = np.take_along_axis(predictions, top_k_indices, axis=1)

        symptoms = []  
        for i in range(predictions.shape[0]):
            # print(f"text {i+1}:")
            for j in range(top_k):
                label_index = top_k_indices[i, j]
                label_name = label_names[label_index]
                prob = top_k_probabilities[i, j]
                # print(f"Label {j+1}: {label_name}, Probability: {prob}")
                
                if prob > 0.3:
                    symptoms.append(label_name)
            
            # print()

        print("Symptoms:")
        print(symptoms)
        
      
        # symptoms = list(decoded_labels)
        document = Document(r"C:\Users\User\WEB\34\web\model\doc_template.docx")

        customer = name
        id = str(age)

        for p in document.paragraphs:
            if "Name : " in p.text:
                p.add_run(customer)
            elif "Age : " in p.text:
                p.add_run(id)

        text_list = [
            '발열', '몸살', '피로감', '소화불량', '가슴통증',
            '두통', '피로감', '신경통', '체중증가', '체중감소',
            '식욕증가', '식욕감퇴', '복통', '심박수증가', '심박수감소',
            '혈압증가', '혈압감소', '변비', '복통', '빈혈',
            '속쓰림', '발진', '어지럼증', '요통', '호흡곤란', '설사']
        
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    for symptom in symptoms:
                        if symptom in text_list and f"{symptom} ☐" in cell_text:
                            cell_text = cell_text.replace(f"{symptom} ☐", f"{symptom} ☑")
                            cell.text = ""
                            cell_paragraph = cell.paragraphs[0]
                            run = cell_paragraph.add_run(f"{symptom} ☑")
                            run.bold = True

        file_name = os.path.splitext(os.path.basename(submit.file.path))[0]
       
        
        new_file_path = 'files/' + file_name + '.docx'
        new_pdf_path = 'files/' + file_name + '.pdf'
        new_image_path1 = 'image/' + file_name +'1'+ '.png'
        new_image_path2 = 'image/' + file_name +'2'+ '.png'
        new_image_path3 = 'image/' + file_name +'3'+ '.png'
        
       # file2 파일 읽기
        data = pd.read_csv(file2)
        # 그래프 그리기
        # plt.plot(df['date'], df['sleep_deep'], label='Sleep Deep')
        # plt.plot(df['date'], df['sleep_rem'], label='Sleep REM')
        # plt.plot(df['date'], df['sleep_light'], label='Sleep Light')
        # plt.xlabel('Date')
        # plt.ylabel('Hours')
        # plt.title('Sleep Patterns')
        # plt.legend()
        # plt.savefig(new_image_path, bbox_inches='tight')
        # plt.show()

        data['date'] = pd.to_datetime(data['date']).dt.strftime('%Y-%m-%d')

        # 도넛 차트 - 최근 7일 평균 수면데이터
        data = data.sort_values('date')
        recent_week = data.tail(7)
        average_data = recent_week[['sleep_deep', 'sleep_rem', 'sleep_light']].mean()

        labels = ['sleep_deep', 'sleep_rem', 'sleep_light']
        sizes = average_data.values.tolist()
        colors = ['#FF9999', '#7BBCFF', '#A2DC85']
        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90, wedgeprops={'edgecolor': 'white'})
        circle = plt.Circle((0, 0), 0.7, color='white')
        fig.gca().add_artist(circle)
        # plt.title('Sleep Distribution - Recent Week', fontweight='bold')
        plt.legend(labels, loc='upper right', bbox_to_anchor=(1.15, 1))
        plt.savefig(new_image_path1, bbox_inches='tight')
        plt.close()

        # 막대 그래프 - 최근 7일 태운 칼로리, 걸음 수
        week_data_one = data[['cal', 'time_activity']].tail(7)
        colors = ['#7BBCFF', '#FFCF8C']
        ax = week_data_one.plot(kind='bar', figsize=(4, 4), color=colors)
        # plt.title("Calories, Steps, and Activity Time for the Last 7 Days")
        plt.xlabel("Date")
        plt.ylabel("Value")
        plt.xticks(rotation=45)
        plt.legend(loc='upper right')
        # 그래프 저장
        plt.savefig(new_image_path2, bbox_inches='tight')
        plt.close()

        # 선 그래프 - 3달치 하루 평균 심박수
        dates = data['date']
        hr_averages = data['hr_average']
        plt.figure(figsize=(8, 4))
        plt.plot(dates, hr_averages, marker='o', color='#7BBCFF', markersize=4)
        plt.xlabel('Date', fontsize=12)
        plt.ylabel('Average Heart Rate', fontsize=12)
        # plt.title('Average Heart Rate by Date', fontsize=14)
        plt.xticks(rotation=45, ha='right', fontsize=8)
        xticks_indices = range(0, len(dates), 7)
        xticks_labels = [date[:10] for date in dates[xticks_indices]]
        plt.gca().set_xticks(dates[xticks_indices])
        plt.gca().set_xticklabels(xticks_labels)
        plt.grid(True, linestyle='--', alpha=0.5)
        # 그래프 저장
        plt.savefig(new_image_path3, bbox_inches='tight')
        plt.close()

        from PIL import Image 
        from docx.shared import Inches      
        # 이미지 파일 불러오기

        image1 = Image.open(new_image_path1)
        image2 = Image.open(new_image_path2)

        # 이미지 크기 가져오기
        width1, height1 = image1.size
        width2, height2 = image2.size

        # 새로운 이미지 생성
        new_width = width1 + width2
        new_height = min(height1, height2)
        new_image = Image.new('RGB', (new_width, new_height))

        # 이미지 합치기(도넛+막대)
        new_image.paste(image1, (0, 0))
        new_image.paste(image2, (width1, 0))

        # 새로운 이미지 저장
        save_path = 'image/' + file_name +'4'+ '.png'
        new_image.save(save_path)

        # 도넛+막대 그래프 삽입
        document.add_paragraph("\n" * 12)
        document.add_picture(save_path, width=Inches(6.2), height=Inches(2.5))

        # 선 그래프 삽입
        document.add_paragraph("\n")
        document.add_picture(new_image_path3, width=Inches(6.5), height=Inches(2.7))
        document.save(new_file_path)

        os.remove(new_image_path1)
        os.remove(new_image_path2)
        
        convert(new_file_path,new_pdf_path)
        
        person = Person(name=name, age=age, file=new_pdf_path, file2=file2)
        person.save()
        return redirect('success')  
    else:
        form = PersonForm()
    return render(request, 'submit.html', {'form': form})

@login_required(login_url='/')
def search_results(request):
    query = request.GET.get('query')
    results = None

    if query:
        results = Person.objects.filter(name__icontains=query)

    context = {'results': results}
    return render(request, 'search_results.html', context)



@login_required(login_url='/')
@never_cache
def search_results(request):
    query = request.GET.get('query')
    results = None

    if query:
        results = Person.objects.filter(name__icontains=query)

    context = {'results': results}
    return render(request, 'search_results.html', context)


def view_docx(request):
    file_path = request.GET.get('file_path', '')  # GET 파라미터로부터 file_path 값 가져오기

    return render(request, 'view_docx.html', {'file_path': file_path})

from django.http import HttpResponse
from django.template.loader import render_to_string
import pdfkit

def pdf_view(request):
    file = request.GET.get('file', '')  # URL 쿼리 매개변수에서 file 값 가져오기

    context = {'file': file}  # 전달할 데이터를 context에 추가
    rendered_html = render_to_string('your_template.html', context)

    # HTML 문자열을 PDF로 변환
    pdf = pdfkit.from_string(rendered_html, False)

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'inline; filename="your_pdf.pdf"'
    response.write(pdf)

    return response


@login_required(login_url='/')
def view_graph(request):
    
    person_name = request.POST.get('person_name')
    f_path = request.POST.get('person_file2')
    person_file2_ = os.path.splitext(os.path.basename(f_path))[0]
    person_file2_path = 'wearable/' + person_file2_ + '.csv'

    print(person_file2_path)
    data = pd.read_csv(person_file2_path)
    
    context = {
        'sleep_deep_day': data['sleep_deep'].iloc[-1]/3600,
        'sleep_rem_day': data['sleep_rem'].iloc[-1]/3600,
        'sleep_light_day': data['sleep_light'].iloc[-1]/3600,
        'sleep_total_day': data['sleep_total'].iloc[-1]/3600,
        
        'sleep_deep_week': data['sleep_deep'].iloc[-7:].sum()/3600,
        'sleep_rem_week': data['sleep_rem'].iloc[-7:].sum()/3600,
        'sleep_light_week': data['sleep_light'].iloc[-7:].sum()/3600,
        'sleep_total_week': data['sleep_total'].iloc[-7:].sum()/3600,
        
        'heart_rate':data['hr_average'].iloc[-1],   
        'heart_rate_week':list(data['hr_average'].iloc[-7:]),   
        
        'cal':data['cal'].iloc[-1], 
        'cal_week':list(data['cal'].iloc[-7:]),
        
        'step':data['step'].iloc[-1],   
        'step_week':list(data['step'].iloc[-7:]),   
        
        'time_activity':data['time_activity'].iloc[-1],   
        'time_activity_week':list(data['time_activity'].iloc[-7:]),
        
        'date': pd.to_datetime(data['date']).dt.strftime('%Y%m%d').tolist()
        #'date' : list(data['date'].iloc[-7:])
    
    }
    
    print(person_name)
    return render(request, 'view_graph.html', {'context':context,'person_name':person_name, 'person_file2': person_file2_path})
    
    



# --------------------------------------------발화자 구분 api-----------------------------------------
class ClovaSpeechClient:
            # Clova Speech invoke URL
                invoke_url = 'https://clovaspeech-gw.ncloud.com/external/v1/5395/9f2a0e330638fb94eafd966c769bd058ca3684fbd91b88c11550b3339cdf0af3'
            # Clova Speech secret key
                secret = '4a481ce97fbd4dcc934a4cfa47a0b0df'

                def req_upload(self, file, completion, callback=None, userdata=None, forbiddens=None, boostings=None,
                            wordAlignment=True, fullText=True, diarization=True):
                    request_body = {
                        'language': 'ko-KR',
                        'completion': completion,
                        'callback': callback,
                        'userdata': userdata,
                        'wordAlignment': wordAlignment,
                        'fullText': fullText,
                        'forbiddens': forbiddens,
                        'boostings': boostings,
                        'diarization': {
                            'enable': True  # 화자 인식 활성화
                        } if diarization else None,
                    }
                    headers = {
                        'Accept': 'application/json;UTF-8',
                        'X-CLOVASPEECH-API-KEY': self.secret
                    }
                    files = {
                        'media': open(file, 'rb'),
                        'params': (None, json.dumps(request_body, ensure_ascii=False).encode('UTF-8'), 'application/json')
                    }
                    response = requests.post(headers=headers, url=self.invoke_url + '/recognizer/upload', files=files)
                    return response
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                
                